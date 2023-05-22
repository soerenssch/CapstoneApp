import pickle
from pathlib import Path
import streamlit as st
from streamlit_option_menu import option_menu
from outscraper import ApiClient
import pandas as pd
import os
import openai
import requests
from tqdm import tqdm
import time
import datetime
import base64
import subprocess
subprocess.run(['pip', 'install', 'openpyxl'])
from docx import Document
from docx.shared import Inches
import io
from io import BytesIO
import matplotlib.pyplot as plt
import numpy
import scipy
from scipy import stats
from datetime import date


st.set_page_config(
    page_title="Sentiment-Analyse",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="expanded",    
)

logo_left_container = st.container()
with logo_left_container:
    st.image("VT_logo.png", use_column_width=False, width=150)

st.markdown(
    """
    <style>
    .footer {
        position: fixed;
        bottom: 0;
        right: 0;
        font-size: 12px;
        padding: 5px;
    }
    </style>
    """,
    unsafe_allow_html=True,
)
st.markdown(
    '<div class="footer">© 2023 Capstone Gruppe Research. All rights reserved.</div>',
    unsafe_allow_html=True,
)


WebScraping = "Web Scraping"
SentimentAnalyse = "Sentiment-Analyse"
MitarbeiterUmfrage = "Auswertung der Mitarbeiterumfrage"
Anleitung = "Beschreibung & Kontakt"

st.sidebar.title("Analysetools:")
input_method = st.sidebar.radio("Wählen Sie eine Option:", (WebScraping, SentimentAnalyse, MitarbeiterUmfrage, Anleitung))


# Depending on which option is selected, display the appropriate information
if input_method == WebScraping:
    with st.sidebar:
        st.header("Beschreibung: Web Scraping")
        st.write("Hier werden Ihnen die Reviews der jeweils ausgewählten Standorte per Web Scraping zum Download zur Verfügung gestellt.")
        # Add any additional information or instructions for this option

elif input_method == SentimentAnalyse:
    with st.sidebar:
        st.header("Beschreibung: Sentiment-Analyse")
        st.write("Hier werden die zuvor gescrapten Reviews hochgeladen und ausgewertet mithilfe einer Sentiment-Analyse ausgewertet.")

elif input_method == MitarbeiterUmfrage:
    with st.sidebar:
        st.header("Beschreibung: Mitarbeiterumfrage")
        st.write("Mithilfe von linearen Regressionen können die Ergebnisse der Mitarbeiterumfrage detailliert ausgewertet werden. Dabei werden die Stärke und Signifikanz der Zusammenhänge ersichtlich und automatisch graphisch dargestellt.")

elif input_method == Anleitung:
    with st.sidebar:
        st.header("Beschreibung")
        st.write("Hier finden Sie einen Überblick über die Funktionsweise der verschiedenen Tools. Zusätzlich liegt ein Kontaktformular vor, sollten Sie auftretende Fragen haben.")


### WebScraping
ID_MAP = {
    "Flint's Praxis für Kleintiere": "ChIJg9LTCjq2kUcRlNmsxEGdNoc",
    'Tierhotel Clinica Alpina Ramosch': "ChIJn9GuzVU3g0cRZAPoW82_WZ0",
    'Tierklinik Clinica Alpina Celerina': "ChIJi2rFwxh9g0cRKh_m8hTN8SA",
    'Tierklinik Clinica Alpina Scuol': "ChIJhf6QYJBHg0cRIe8DYQz8JV8",
    'Kleintierklinik am Damm, Gossau': "ChIJsXk3JcvgmkcRmMwJslg2xcU",
    'Kleintierpraxis Aabach Uster': "ChIJE9cxb6CkmkcRH3ybj-adIbE",
    'Kleintierpraxis Au': "ChIJm8-ANlQRm0cRnd3hTS279ag",
    'Kleintierpraxis Bachmatt Eschenbach': "ChIJnU4TOnT9j0cR6HQvLSgsS5Q",
    'Kleintierpraxis Basel Central': "ChIJAR4_G0-4kUcRQAgkjVkUSek",
    'Kleintierpraxis Basel Spalen': "ChIJqRyBUAG5kUcRVcVH76IXaFE",
    'Kleintierpraxis Bülach': "ChIJS9aRbdJ1kEcRhKrcNtRWsuA",
    'Kleintierpraxis Fällanden': "ChIJzZH-veOjmkcR7GUzmVWaezY",
    'Kleintierpraxis Glattbrugg': "ChIJ9UNysAigmkcR5QgQSw4lpiY",
    'Kleintierpraxis Gümligen': "ChIJv5D40wE3jkcRFKWuw1rHTtE",
    'Kleintierpraxis im Bahnhof, Aathal': "ChIJ6Z-zRnG7mkcRPBbZp4n8iww",
    'Kleintierpraxis Küssnacht am Rigi': "ChIJHY_A-m7_j0cRHUpOOIwa8Yo",
    'Kleintierpraxis Mühlebach Oftringen': "ChIJf_mbzJQvkEcRicJ5NcgPDuE",
    'Kleintierpraxis Münchenstein': "ChIJvwTRBxe4kUcRfEUl0F_UF4Y",
    'Kleintierpraxis Muri': "ChIJ5-00QTwFkEcRG9Hk9XEzIM8",
    'Kleintierpraxis Oensingen': "ChIJg3c2oE7SkUcRy_Yz-zeQUww",
    'Kleintierpraxis Regensdorf': "ChIJF5STJ2cLkEcRDRye_BfHcDc",
    'Kleintierpraxis Schlieren': "ChIJERk9VbMNkEcRPy81sQi1wic",
    'Kleintierpraxis St. Gallen Ost': "ChIJTeRhD_Qem0cRdZHi-gV_KQc",
    'Kleintierpraxis Stansstad': "ChIJO7t-Rkb3j0cRhQ4zkf4JuBs",
    'Kleintierpraxis Telli Aarau': "ChIJaZqDHPw7kEcR-Fp5zfYNNdI",
    'Kleintierpraxis Therwil': "ChIJdUCYp3bHkUcRnqgwh2XgkjI",
    'Kleintierpraxis Trimbach': "ChIJlxwBXNMxkEcRVFq9hfD4rYc",
    'Kleintierpraxis Turbenthal': "ChIJGyWX2Q2WmkcRx1gLsoFhMz8",
    'Kleintierpraxis Wettingen': "ChIJGY-uTiltkEcR5qPVT9EHlVU",
    'Kleintierpraxis Winterthur im Zentrum': "ChIJK9ieJXmZmkcRHtJtVS47ZwE",
    'Kleintierpraxis Zug': "ChIJtSFUYUWqmkcRnQZS-2NJx-w",
    'Kleintierpraxis Zuzwil': "ChIJPWi96nfpmkcRMxRLWqMgC5s",
    'Tierklinik Basel': "ChIJpSNsIru5kUcRx-XRYVJBeU4",
    'Tierklinik Nesslau': "ChIJb_dA_kPXmkcRAGqsVT1wLFE",
    'Tierklinik Oberland Pfäffikon': "ChIJkTyE5WG8mkcR9FjBoEAKIEE",
    'Tierklinik Oberland Saland': "ChIJ2RDiJW6-mkcRRKSDQ0onkc4",
    'Tierklinik Zürich': "ChIJNXoM5QOhmkcRBNaok_HRShE",
    'Zentrum für Tiermedizin Klettgau': "ChIJH2Ls-3Z8kEcRvq0oZ5ctkrk",
}

input_Outscraper = []
checkbox_states = {label: False for label in ID_MAP.keys()}
checkbox_states["Select all"] = False

def update_selection(ID_MAP, checkbox_states, key):
    checkbox_states[key] = not checkbox_states[key]  # toggle the state of the clicked checkbox
    if key == "Select all":
        for label in ID_MAP.keys():
            checkbox_states[label] = checkbox_states[key]
    else:
        if not checkbox_states[key]:
            checkbox_states["Select all"] = False  # if any checkbox is unchecked, uncheck "Select all"
        else:
            checkbox_states["Select all"] = all(checkbox_states.values())  # if all checkboxes are checked, check "Select all"
    return checkbox_states

if input_method == WebScraping:
    st.title("Web Scraping")
    st.write("Mithilfe von Web Scraping werden ab einem ausgewählten Zeitpunkt alle Google Reviews der gewählten Standorte zum Download bereitgestellt. Bitte wählen Sie die Standorte aus, die Sie interessieren.")


    select_all = st.checkbox("Alle Standorte auswählen")

    place_checkboxes = {}
    for place in ID_MAP:
        if select_all:
            place_checkboxes[place] = st.checkbox(place, value=True)
            input_Outscraper.append(ID_MAP[place])
        else:
            place_checkboxes[place] = st.checkbox(place)
        
        if place_checkboxes[place]:
            place_id = ID_MAP[place]
            if place_id not in input_Outscraper:
                input_Outscraper.append(place_id)
        else:
            place_id = ID_MAP[place]
            if place_id in input_Outscraper:
                input_Outscraper.remove(place_id)

    
    date_input = st.date_input('Geben Sie das Datum an, ab dem die Reviews exportiert werden sollen:')
    if date_input > datetime.datetime.today().date():
        st.error('Fehler: Datum darf nicht in der Zukunft liegen!')
    else:
        year = date_input.year
        month = date_input.month
        day = date_input.day

    my_date = datetime.datetime(year, month, day)
    timestamp = my_date.timestamp()
    timestamp = int(timestamp)
    st.session_state.timestamp = timestamp

    st.write("Melden Sie sich mit dem nachfolgenden Link bei Outscraper an, um Ihren eigenen API-Key zu erstellen: https://outscraper.com/")
    
    Outscraper_APIKey = st.text_input("Geben Sie hier Ihren Outscraper API-Key an:")
    client = ApiClient(api_key=Outscraper_APIKey)

    

    submit = st.button("Web Scraping durchführen")
    if submit: 
        st.write("Web Scraping wird durchgeführt!")
        @st.cache_data(ttl=600)
        def scrape_google_reviews(query, timestamp):
            results = client.google_maps_reviews([query], sort='newest', cutoff=timestamp, reviews_limit=1000, language='de')
            return results


        results = scrape_google_reviews(input_Outscraper, timestamp)

        data = []
        for place in results:
            name = place['name']
            for review in place.get('reviews_data', []):
                review_text = review['review_text']
                review_rating = review['review_rating']
                date = review["review_datetime_utc"]
                data.append({'Standort': name, 'Review': review_text, 'Rating': review_rating, 'Datum': date})
        df = pd.DataFrame(data)
        st.dataframe(df)

        

        def generate_csv(df):
            return df.to_csv(index=False)

        if st.download_button(label='Download CSV', data=generate_csv(df), file_name='data.csv', mime='text/csv'):
            pass

### SentimentAnalyse  
if input_method == SentimentAnalyse:
    

    def download_file(file):
        b = BytesIO()
        file.savefig(b, format='png')
        b.seek(0)
        return b

    # Main code
    st.title("Sentiment-Analyse")
    file = st.file_uploader("Laden Sie Ihre Datei hier hoch:", type=["csv"])
    if file is not None:
        df = pd.read_csv(file)
        df = df.dropna()

        if "Datum" in df.columns:

            # add checkbox to allow manual date selection
            manual_date_selection = st.checkbox("Den Analysezeitraum manuell festlegen")
            df["Datum"] = pd.to_datetime(df["Datum"])
            if manual_date_selection:

                
                # start_date = st.date_input("Start date")
                # end_date = st.date_input("End date")
                start_date = st.date_input("Start der Auswertung:", value=pd.to_datetime(df['Datum']).min().date())
                end_date = st.date_input("Ende der Auswertung:", value=pd.to_datetime(df['Datum']).max().date())

                # mask = (df["date"] >= start_date) & (df["date"] <= end_date)
                mask = (pd.to_datetime(df['Datum']).dt.date >= start_date) & (pd.to_datetime(df['Datum']).dt.date <= end_date)
                df = df.loc[mask]

        st.dataframe(df)

        if "Rating" in df.columns:
                
            
            def generate_plot(df):
                fig, ax = plt.subplots(figsize=(12, 8))
                counts, bins, patches = ax.hist(df['Rating'], bins=5, color='#132f55', edgecolor='white')
                ax.set_xlabel('Bewertung')
                ax.set_ylabel('Anzahl')
                ax.set_title('Verteilung der Bewertungen')
                tick_labels = ['1 Stern', '2 Sterne', '3 Sterne', '4 Sterne', '5 Sterne']
                tick_positions = [1.415, 2.25, 3, 3.85, 4.63]
                ax.set_xticks(tick_positions)
                ax.set_xticklabels(tick_labels, ha='center')

                df_Anzahl = df['Rating'].count()
                df_mean = df['Rating'].mean()
                df_std = df['Rating'].std()
                ax.legend([f'Anzahl Bewertungen: {df_Anzahl:.0f}\n\n Durchschnitt: {df_mean:.2f}\n\n Standardabweichung: {df_std:.2f}'])

                # Add numbers above bars
                for count, patch in zip(counts, patches):
                    height = patch.get_height()
                    ax.annotate(f'{count:.0f}', xy=(patch.get_x() + patch.get_width() / 2, height), xytext=(0, 3),
                                textcoords='offset points', ha='center', va='bottom')

                return fig

        
            fig = generate_plot(df)
            st.pyplot(fig)            

            download = download_file(fig)
            st.download_button(
                label='Download Grafik',
                data=download,
                file_name='grafik.png',
                mime='image/png'
            )
            
        Spalte = st.text_input("Wie heisst die Spalte, die Sie auswerten möchten?")
        if not Spalte:
            Spalte = "Review"

        st.write("Melden Sie sich mit dem nachfolgenden Link bei OpenAI an, um Ihren API-Key zu erstellen: https://platform.openai.com/")
        OpenAI_API = st.text_input("Geben Sie hier Ihren OpenAI API-Key an:")
        
        openai.api_key = OpenAI_API
        GPT_API_URL = "https://platform.openai.com/"
        all_reviews = "\n".join(df[Spalte].tolist())
        if st.button("Sentiment-Analyse starten"):

            @st.cache_data(ttl=600)
            def generate_proscons_list(text):
                word_blocks = text.split(' ')
                block_size = 1700
                blocks = [' '.join(word_blocks[i:i + block_size]) for i in range(0, len(word_blocks), block_size)]

                proscons = []

                for block in tqdm(blocks, desc="Processing blocks", unit="block"):
                    messages = [
                        {"role": "system", "content": "Du bist ein KI-Sprachmodell, das darauf trainiert ist, eine Liste der häufigsten Stärken und Schwächen einer Tierarztpraxis auf der Grundlage von Google Bewertungen zu erstellen. Es ist dir nicht erlaubt, über die Liste der Stärken und Schwächen hinaus einen Output zu generieren, also keine Zusammenfassung der Fragestellung oder des Ergebnisses."},
                        {"role": "user", "content": f"Erstelle auf der Grundlage der folgenden Google-Bewertungen eine Liste mit den häufigsten Stärken und Schwächen der Tierarztpraxis: {block}"}
                    ]

                    completion = openai.ChatCompletion.create(
                        model="gpt-3.5-turbo",
                        messages=messages,
                        # You can change the max_tokens amount to increase or decrease the length of the results pros and cons list. If you increase it too much, you will exceed chatGPT's limits though.
                        max_tokens=250,
                        n=1,
                        stop=None,
                            # You can adjust how "creative" (i.e. true to the original reviewer's intent) chatGPT will be with it's summary be adjusting this temperature value. 0.7 is usually a safe amount
                        temperature=0.7
                    )

                    procon = completion.choices[0].message.content
                    proscons.append(procon)

                    # Gefundene Stärken und Schwächen in einer Liste zusammenfassen 
                combined_proscons = "\n\n".join(proscons)
                return combined_proscons
            summary_proscons = generate_proscons_list(all_reviews)

            df_proscons = pd.DataFrame()
            list_proscons = []
            list_proscons.append(summary_proscons)
            df_proscons["pros_cons"] = list_proscons

            # Create Word document
            document = Document()
            document.add_heading("Stärken und Schwächen Zusammenfassung", level=0)
            table = document.add_table(rows=len(df_proscons.index), cols=1)
            for i, row in df_proscons.iterrows():
                table.cell(i, 0).text = row["pros_cons"]
            document.add_page_break()

            # Download Word document
            with io.BytesIO() as output:
                document.save(output)
                if st.download_button(label='Download', data=output.getvalue(), file_name='Ergebnisse.docx', mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'):
                    pass

if input_method == MitarbeiterUmfrage:
    st.title("Auswertung der Mitarbeiterumfrage")

    file = st.file_uploader("Laden Sie Ihre Datei hier hoch:", type=["csv"])
    if file is not None:
        df = pd.read_csv(file)
        
        def cleaning(df):
            df = df.drop(df.columns[[0,1,15]], axis=1 )
            df.columns.values[0] = "Wertschätzung"
            df.columns.values[1] = "Team"
            df.columns.values[2] = "Ausgeglichenheit"
            df.columns.values[3] = "Zufriedenheit als Mitarbeiter"
            df.columns.values[4] = "Weiterempfehlung als Arbeitgeber"
            df.columns.values[5] = "Zufriedenheit der Kunden"
            df.columns.values[6] = "Weiterempfehlung als Tierarztpraxis/-klinik"
            df.columns.values[7] = "Kennen der Werte"
            df.columns.values[8] = "Identifikation mit Werten"
            df.columns.values[9] = "Anweisungen"
            df.columns.values[10] = "Image in der CH"
            df.columns.values[11] = "Kommunikation Head Office"
            df.columns.values[12] = "Weiterbildungen und Karriere"
            return df
        
        

        st.dataframe(df)

        if st.button("Auswertung starten"):
            df = cleaning(df)
            def regressionen(df):
                p_total = pd.DataFrame()
                r_total = pd.DataFrame()
                complete = pd.DataFrame()
                today = date.today()
                for column_x in df:
                    p_dict = {}
                    r_dict = {}
                    for column_y in df:
                        df_copy = df.dropna(subset=[column_x, column_y])
                        x = df_copy[column_x]
                        y = df_copy[column_y]  
                        y = y.dropna()
                        slope, intercept, r, p, std_err = stats.linregress(x, y)    
                        p_dict['{} vs. {}'.format(column_x, column_y)] = p
                        r_dict['{} vs. {}'.format(column_x, column_y)] = r 
                    p_df = pd.DataFrame.from_dict(p_dict, orient='index')
                    r_df = pd.DataFrame.from_dict(r_dict, orient='index')
                    p_total = pd.concat([p_total, p_df])
                    r_total = pd.concat([r_total, r_df])
                r_total = r_total.reset_index()
                p_total = p_total.reset_index()
                r_total = r_total.rename(columns={"index": "Parameter", 0: "r"})
                p_total = p_total.rename(columns={"index": "Parameter", 0: "p"})
                p_total= p_total.round(6)
                r_total= r_total.round(6)
                r_total = r_total.drop_duplicates(subset=["r"], keep="first")
                complete = pd.merge(r_total, p_total, on="Parameter", how="left")
                complete = complete[complete.r != 1]
                complete=complete.sort_values(by=['r'], ascending=False)
                complete = complete.reset_index()
                complete = complete.drop(columns="index")
                signifikant = complete[complete.p <= 0.05]
                return complete, signifikant
            
            complete, signifikant = regressionen(df)

            st.dataframe(signifikant)

            def generate_csv(signifikant):
                return signifikant.to_csv(index=False)

            if st.download_button(label='Download CSV', data=generate_csv(signifikant), file_name='Signifikanteste_Ergebnisse.csv', mime='text/csv'):
                pass



            def create_plots(column_x, column_y, df):
                x = df[column_x]
                y = df[column_y] 
                slope, intercept, r, p, std_err = stats.linregress(x, y)    
                def myfunc(x):
                    return slope * x + intercept
                mymodel = list(map(myfunc, x))
                plt.ylim([0, 5])
                plt.xlim([0, 5])
                plt.scatter(x, y, color=("#132f55"))
                plt.plot(x, mymodel,color=("#d52f89"))
                plt.title("{} vs. {}".format(column_x, column_y))
                plt.xlabel("{}".format(column_x))
                plt.ylabel("{}".format(column_y))
                plt.show()

            r = signifikant["r"]
            p = signifikant["p"] 
            list_r = []
            list_p = []
            for i in range(len(p)):
                list_p.append(p[i])
                list_r.append(r[i])

            parameter = signifikant["Parameter"]

            for i in range(len(signifikant)):
                x_axis = parameter[i].split(' vs. ')[0]
                y_axis = parameter[i].split(' vs. ')[1]
                fig, ax = plt.subplots(figsize=(5, 3))
                ax.scatter(df[x_axis], df[y_axis], color=("#132f55"))
                ax.set_xlabel(x_axis)
                ax.set_ylabel(y_axis)
                ax.set_title("{} vs. {}".format(x_axis, y_axis))
                x = df[x_axis]    
                y = df[y_axis] 
                slope, intercept, r, p, std_err = stats.linregress(x, y)    
                def myfunc(x):
                    return slope * x + intercept
                mymodel = list(map(myfunc, x))
                
                ax.plot(x, mymodel, color=("#d52f89"), label=f"p={list_p[i]:.4f} \n\n r={list_r[i]:.4f}")
                ax.legend()
                
                def download_plot(plot):
                    output_buffer = io.BytesIO()
                    plot.savefig(output_buffer, format='png')
                    output_buffer.seek(0)
                    return output_buffer
                
                st.pyplot(fig)
                buffer = download_plot(fig)
                st.download_button(
                    label=f"Download Grafik {x_axis} vs. {y_axis}",
                    data=buffer,
                    file_name="Grafik_{}_{}.png".format(x_axis, y_axis),
                    mime='image/png'
                )
                


                

if input_method == Anleitung:
    st.title("Beschreibung & Kontakt")
    st.write("""Um das Tool optimal nutzen zu können, müssen Sie sich bei Outscraper und OpenAI einen Account anlegen und einen API-Key erstellen. WICHTIG: Geben Sie Ihren Key nicht an Dritte weiter! 
    Die Links zur Anmeldung finden Sie hier:

    \n\nOutscraper: https://outscraper.com/

    \n\nOpenAI: https://platform.openai.com/\n\n


    \n\n\n\nWeb Scraping:

    \n\nBeim sogenannten Web Scraping greift das Programm auf die Google-Maps Bewertungen der VetTrust zu und fasst diese innerhalb einer „.csv (comma-seperated-values)“-Datei zusammen, welche zentral über das Programm Excel geöffnet werden kann. Um die relevanten Daten zu erhalten, muss zudem angegeben werden, von welchem Standort und ab welchem Zeitpunkt die Exporte benötigt werden. Darüber hinaus wird der zuvor erstellte API-Key benötigt, um das Web Scraping durchzuführen. Die Kosten liegen bei 0,002 Dollar pro Review, d.h. bei 200 Reviews kostet die Auswertung 0,04 Dollar. Je nach Anzahl der Standorte sowie dem ausgewählten Zeitraum dauert diese Applikation wenige Sekunden bis einige Minuten. Das Ergebnis kann abschliessend heruntergeladen werden, um es entweder manuell zu betrachten oder mit der Sentiment-Analyse näher zu analysieren.

    \n\n\n\nSentiment-Analyse:

    \n\nHier wird die beim Web Scraping erstellte .csv-Datei, die nun ausgewertet werden soll, zur weiteren Auswertung wieder hochgeladen. Das Modell analysiert nun die Texte auf positive und negative Aspekte, weshalb die Sentiment-Analyse auch auf andere Datensätze mit gleichem Format angewendet werden kann. Zu beachten ist, dass alle Textausschnitte in derselben Spalte vorliegen müssen, da die Analyse nur für eine Spalte möglich ist. Um die Analyse durchzuführen, muss nach Upload der Daten zusätzlich der Name der Spalte angegeben werden, die nachfolgend ausgewertet werden soll. Die auszuwertende Spalte, der im Web Scraping exportierten Daten ist immer durch den Namen „Review“ gekennzeichnet - jedoch kann dies bei eigenen Datensätzen abweichen. In diesem Fall muss der Name der auszuwertenden Spalte manuell eingetragen werden. In einem letzten Schritt muss auch hier wieder der passende API-Key eingefügt werden. Die Kosten liegen dabei bei 0,002 Dollar je 1000 Tokens, wobei 1000 Tokens ca. 750 Wörter sind. Es zählen ausserdem nicht nur die Wörter des Inputs, sondern auch die des Outputs. D.h. bei 200 Reviews à 20 Wörtern als Input und der gleichen Anzahl als Output, entstehen Kosten von ca. 0,02 Dollar. In Abhängigkeit der Grösse des jeweiligen Datensatzes kann die Analyse einige Sekunden bis wenige Minuten dauern. Das endgültige Ergebnis kann, nach Abschluss der Sentiment-Analyse, als Word-Datei heruntergeladen werden. In dieser sind die Stärken und Schwächen respektive die positiven und negativen Aspekte der Bewertungen übersichtlich aufgelistet.
    
    \n\nAuswertung der Mitarbeiterumfrage:

    \n\nZuerst müssen Sie die .csv-Datei der Mitarbeiterumfrage hochladen. Danach müssen Sie nur noch auf “Auswertung starten” klicken. Unser Modell vergleicht die Antworten jeder einzelnen Frage miteinander und findet mithilfe von linearen Regressionen Zusammenhänge in den Antworten. Sie erhalten eine neue .csv-Datei fertig zum Herunterladen, die eine Tabelle mit den Zusammenhängen und dem jeweiligen r-Wert und p-Wert darstellt. 
    \n\nDer r-Wert ist ein Indikator für die Korrelation von der Ergebnisse von zwei Fragen. Je näher r bei Null liegt, desto schwächer ist der lineare Zusammenhang. Positive r-Werte zeigen eine positive Korrelation an, bei der die Werte beider Variable tendenziell gemeinsam ansteigen. Negative r-Werte zeigen eine negative Korrelation an, bei der die Werte einer Variable tendenziell ansteigen, wenn die Werte der anderen Variable fallen. 
    \n\nDer p-Wert misst die Wahrscheinlichkeit, dass ein in der Umfrage beobachteter Unterschied zwischen zwei Fragen zufällig entstanden sein könnte. Ist diese Wahrscheinlichkeit gering, dann ist der beobachtete Unterschied vermutlich statistisch signifikant und eventuell auch auf die Antworten von Mitarbeitenden übertragbar, die nicht bei der Umfrage teilgenommen haben. Das Signifikanzniveau wurde auf 0.05 eingestellt, deshalb werden nur Zusammenhänge angezeigt, welche einen p-Wert von unter 0.05 erreicht haben und somit statistisch signifikant sind. 
    \n\nUm die letzte Spalte der Mitarbeiterumfrage ("Sonstiges Feedback") auszuwerten, empfehlen wir Ihnen die Nutzung der Sentiment-Analyse, wie im oberen Schritt beschrieben.
    """)

    st.header("Kontaktformular")
    st.write("Bei Fragen können Sie uns hier gerne kontaktieren!")

    contact_form = """
    <form action="https://formsubmit.co/soeren.schlisske@web.de" method="POST">
        <input type="text" name="name" placeholder="Ihr Name" required>
        <input type="email" name="email" placeholder="Ihre Email-Adresse" required>
        <textarea name="message" placeholder="Ihre Nachricht"></textarea>
        <button type="submit">Senden</button>
    </form>
    """

    st.markdown(contact_form, unsafe_allow_html=True)

    def local_css(file_name):
        with open(file_name) as f:
            st.markdown(f'<style>{f.read()}</style>', unsafe_allow_html=True)

    local_css("style/style.css")


   